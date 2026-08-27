import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

HEX_TEAL = "0F7E9B"
HEX_AMBER = "D67B19"
HEX_LIGHT_BG = "E6F4F8"
HEX_LIGHT_AMBER = "FDF3E8"
HEX_BORDER = "CBD5E1"
HEX_DARK = "1E293B"

RGB_TEAL = RGBColor(15, 126, 155)
RGB_AMBER = RGBColor(214, 123, 25)
RGB_DARK = RGBColor(30, 41, 59)
RGB_MUTED = RGBColor(100, 116, 139)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGB_TEAL
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGB_AMBER
    return p

def format_doc_header(doc, title, subtitle):
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("THE THINKING EXPERIMENT | CVPM MODULE")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(8.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGB_AMBER

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_TEAL

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(10)
    run_desc = p_desc.add_run(subtitle)
    run_desc.font.name = 'Arial'
    run_desc.font.size = Pt(9.5)
    run_desc.font.color.rgb = RGB_MUTED

def add_callout(doc, title, bullets, is_amber=False):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.8)
    cell = tbl.cell(0, 0)
    set_cell_background(cell, HEX_LIGHT_AMBER if is_amber else HEX_LIGHT_BG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run_t = p.add_run(title)
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10)
    run_t.font.bold = True
    run_t.font.color.rgb = RGB_AMBER if is_amber else RGB_TEAL
    
    for b in bullets:
        pb = cell.add_paragraph()
        pb.paragraph_format.space_after = Pt(2)
        run_b = pb.add_run(b)
        run_b.font.name = 'Arial'
        run_b.font.size = Pt(9)
        run_b.font.color.rgb = RGB_DARK
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)

def generate_cv_student_handout(filepath):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
        
    format_doc_header(doc, "Constant Velocity Motion & Dual-Graph Translation", "Student Inquiry & Quantitative Graphing Practice")
    
    add_callout(doc, "📐 Fundamental Constant Velocity Mathematical Tools", [
        "• Slope of Position-Time (x-t) Graph = Velocity (v):\n   v = Δx / Δt = (x_final - x_initial) / (t_final - t_initial)",
        "• Area under Velocity-Time (v-t) Graph = Displacement (Δx):\n   Δx = v · Δt   →   x_final = x_initial + v · Δt",
        "• Pawprint Motion Map: Dropped every Δt = 0.28 s. Even spacing = constant speed; vertical stack = stationary at rest.",
        "• Coordinate System: + = Moving Right | - = Moving Left | Origin = 0.0 m"
    ])
    
    add_heading_1(doc, "📝 Part 1: Quantitative Dual-Graph Construction Tasks")
    p_intro = doc.add_paragraph("For each problem, analyze the given numerical graph, perform the required slope/area calculations, construct the complementary numerical graph, sketch the pawprint motion map, and write a complete physical description of Coco's run.")
    p_intro.paragraph_format.space_after = Pt(8)
    
    problems = [
        {
            "num": "Problem 1: Given Position-Time Graph → Construct Velocity-Time Graph",
            "given_type": "Position-Time (x-t) Given",
            "given_data": "Coco starts at x₀ = 2.0 m at t = 0 s and trots steadily to x = 14.0 m at t = 4.0 s.",
            "calc_prompt": "Step 1: Calculate Coco's constant velocity (v = Δx / Δt):\n   v = (14.0 m - 2.0 m) / (4.0 s - 0 s) = __________ m/s",
            "graph_prompt": "Step 2: Draw the numerical Velocity-Time (v-t) graph from t = 0 to 4.0 s.\n   (Label the horizontal line at v = _____ m/s with axes: t from 0 to 4s, v from -5 to +5 m/s)",
            "map_prompt": "Step 3: Sketch Coco's pawprint trail on the runway (0m to 20m).",
            "desc_prompt": "Step 4: Written Narrative: Describe Coco's motion including starting position, speed, and direction."
        },
        {
            "num": "Problem 2: Given Velocity-Time Graph → Construct Position-Time Graph",
            "given_type": "Velocity-Time (v-t) Given",
            "given_data": "Coco maintains a constant velocity v = -3.0 m/s from t = 0 s to t = 4.0 s. Her initial position is x₀ = 18.0 m.",
            "calc_prompt": "Step 1: Calculate Coco's displacement (Δx = v · Δt) and final position (x_f = x₀ + Δx):\n   Δx = (-3.0 m/s) · (4.0 s) = __________ m\n   x_final = 18.0 m + (_____) = __________ m",
            "graph_prompt": "Step 2: Draw the numerical Position-Time (x-t) graph from t = 0 to 4.0 s.\n   (Plot points (0 s, 18.0 m) and (4.0 s, _____ m), connect with a straight line)",
            "map_prompt": "Step 3: Sketch Coco's pawprint trail showing direction of motion.",
            "desc_prompt": "Step 4: Written Narrative: Why is the slope negative? What does the negative sign physically mean?"
        },
        {
            "num": "Problem 3: Multi-Stage Motion (Trot → Sniff at Rest → Trot)",
            "given_type": "Position-Time (x-t) Given",
            "given_data": "Stage 1 (0 to 2 s): Moves from x = 0 m to x = 8.0 m.\nStage 2 (2 to 4 s): Remains at x = 8.0 m.\nStage 3 (4 to 6 s): Moves from x = 8.0 m to x = 14.0 m.",
            "calc_prompt": "Step 1: Calculate the velocity for each stage:\n   • Stage 1 (0 to 2s): v₁ = (8.0 - 0) / (2 - 0) = __________ m/s\n   • Stage 2 (2 to 4s): v₂ = (8.0 - 8.0) / (4 - 2) = __________ m/s\n   • Stage 3 (4 to 6s): v₃ = (14.0 - 8.0) / (6 - 4) = __________ m/s",
            "graph_prompt": "Step 2: Construct the 3-step piecewise Velocity-Time (v-t) graph from t = 0 to 6.0 s.",
            "map_prompt": "Step 3: Sketch Coco's pawprint trail (show stacking where stationary).",
            "desc_prompt": "Step 4: Written Narrative: Compare Coco's speed in Stage 1 vs Stage 3. Which segment has the steeper slope?"
        },
        {
            "num": "Problem 4: Turnaround Motion (Fetch & Return)",
            "given_type": "Velocity-Time (v-t) Given",
            "given_data": "Stage 1 (0 to 2 s): v = +4.0 m/s (running for ball)\nStage 2 (2 to 3 s): v = 0.0 m/s (grabbing ball at rest)\nStage 3 (3 to 5 s): v = -4.0 m/s (returning with ball)\nInitial Position: x₀ = 2.0 m",
            "calc_prompt": "Step 1: Calculate displacements and positions at key time intervals:\n   • t = 2 s: Δx₁ = (+4.0 m/s)(2.0 s) = +8.0 m   →   x(2s) = 2.0 + 8.0 = 10.0 m\n   • t = 3 s: Δx₂ = (0 m/s)(1.0 s) = 0.0 m       →   x(3s) = 10.0 m\n   • t = 5 s: Δx₃ = (-4.0 m/s)(2.0 s) = -8.0 m   →   x(5s) = 10.0 - 8.0 = __________ m",
            "graph_prompt": "Step 2: Construct the exact numerical Position-Time (x-t) graph from t = 0 to 5.0 s.",
            "map_prompt": "Step 3: Sketch the pawprint trail showing the turnaround point.",
            "desc_prompt": "Step 4: Written Narrative: What is Coco's total distance traveled vs. her net displacement?"
        }
    ]
    
    for prob in problems:
        add_heading_2(doc, prob["num"])
        
        # Grid Card Table
        tbl = doc.add_table(rows=4, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.columns[0].width = Inches(3.4)
        tbl.columns[1].width = Inches(3.4)
        
        # Row 0: Given info & Calculations
        c00 = tbl.cell(0, 0)
        set_cell_background(c00, "F8FAFC")
        set_cell_margins(c00, top=80, bottom=80, left=100, right=100)
        p00 = c00.paragraphs[0]
        p00.add_run("📋 " + prob["given_type"] + "\n").bold = True
        p00.add_run(prob["given_data"]).font.size = Pt(8.5)
        
        c01 = tbl.cell(0, 1)
        set_cell_background(c01, HEX_LIGHT_BG)
        set_cell_margins(c01, top=80, bottom=80, left=100, right=100)
        p01 = c01.paragraphs[0]
        p01.add_run("🔢 Quantitative Calculations\n").bold = True
        p01.add_run(prob["calc_prompt"]).font.size = Pt(8.5)
        
        # Row 1: Graph grid drawing box
        c10 = tbl.cell(1, 0)
        set_cell_margins(c10, top=80, bottom=80, left=100, right=100)
        p10 = c10.paragraphs[0]
        p10.add_run("📈 Constructed Complementary Graph:\n").bold = True
        p10.add_run(prob["graph_prompt"] + "\n\n\n\n\n\n[ Graph Grid Space: Draw Axes & Label Numerical Points ]").font.size = Pt(8.5)
        
        c11 = tbl.cell(1, 1)
        set_cell_margins(c11, top=80, bottom=80, left=100, right=100)
        p11 = c11.paragraphs[0]
        p11.add_run("🐾 Pawprint Motion Map Sketch:\n").bold = True
        p11.add_run(prob["map_prompt"] + "\n\n\n0m ───────────────────────────── 20m\n[ Draw Pawprints & Stack where at Rest ]").font.size = Pt(8.5)
        
        # Row 2 & 3: Written Narrative
        c2 = tbl.cell(2, 0)
        c2.merge(tbl.cell(2, 1))
        set_cell_background(c2, HEX_LIGHT_AMBER)
        set_cell_margins(c2, top=80, bottom=80, left=100, right=100)
        p2 = c2.paragraphs[0]
        p2.add_run("💬 Motion Description & Physical Interpretation\n").bold = True
        p2.add_run(prob["desc_prompt"] + "\n\n\n").font.size = Pt(8.5)
        
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(6)
        
    add_heading_1(doc, "🧠 Part 2: Synthesis & Self-Assessment Check")
    synth_questions = [
        "1. Graphical Meaning of Slope: A student claims that 'a horizontal line on any graph means the object is stopped.' Explain why this is true for Position-Time graphs, but completely false for Velocity-Time graphs.",
        "2. Displacement vs. Distance: In Problem 4, calculate Coco's Total Distance Traveled (scalar sum) and her Net Overall Displacement (vector change in position Δx = x_final - x_initial). Explain why they are different.",
        "3. Simulation Verification: Open the Constant Velocity (CVPM) tab in the simulator. Test CV-1 through CV-8. Which scenario in the simulator matches Problem 4?"
    ]
    
    for q in synth_questions:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        run_q = p_q.add_run(q)
        run_q.font.name = 'Arial'
        run_q.font.size = Pt(9.5)
        run_q.font.color.rgb = RGB_DARK
        
        p_ans = doc.add_paragraph("Answer:\n\n\n")
        p_ans.paragraph_format.space_after = Pt(6)
        p_ans.runs[0].font.color.rgb = RGB_MUTED
        p_ans.runs[0].font.size = Pt(9)

    doc.save(filepath)
    print(f"Generated CV Student Handout: {filepath}")

def generate_cv_teacher_guide(filepath):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
        
    format_doc_header(doc, "Constant Velocity Motion & Dual-Graph Translation — Teacher Guide", "Complete Solutions, Step-by-Step Calculations & PER Diagnostic Notes")
    
    add_callout(doc, "📋 Pedagogical Scope & Unit Design (CVPM)", [
        "• Standalone Unit Alignment: This worksheet supports on-level and honors physics units on 1D Uniform Motion (Constant Velocity Particle Model).",
        "• Dual-Graph Fluency: Requires bidirectional translation (x-t → v-t and v-t → x-t) using slope and area operations with explicit numerical values.",
        "• Physical Interpretation: Connects graphical slope/area to real-world dog motion, directional pawprint spacing, and starting position offsets."
    ])
    
    add_heading_1(doc, "🔑 Complete Step-by-Step Answer Key & Numerical Solutions")
    
    solutions = [
        {
            "title": "Problem 1: Given x-t → Construct v-t",
            "calc": "v = (14.0 m - 2.0 m) / (4.0 s - 0 s) = +12.0 m / 4.0 s = +3.0 m/s.",
            "graph": "Horizontal straight line at v = +3.0 m/s extending from t = 0 to t = 4.0 s.",
            "map": "Pawprints start at x = 2.0 m and are evenly spaced to the right ending at x = 14.0 m.",
            "narrative": "Coco begins 2.0 m to the right of the origin and jogs steadily to the right at a constant speed of 3.0 m/s for 4.0 seconds.",
            "misconception": "Students may mistakenly use v = 14/4 = 3.5 m/s by forgetting to subtract the non-zero initial position x₀ = 2.0 m."
        },
        {
            "title": "Problem 2: Given v-t → Construct x-t",
            "calc": "Δx = (-3.0 m/s)(4.0 s) = -12.0 m. Final Position: x_f = 18.0 m + (-12.0 m) = +6.0 m.",
            "graph": "Straight line starting at (0 s, 18.0 m) sloping downward with steady slope -3.0 m/s to (4.0 s, 6.0 m).",
            "map": "Pawprints start at x = 18.0 m and point left, spaced evenly ending at x = 6.0 m.",
            "narrative": "Coco starts at x = 18.0 m and walks to the left at a steady speed of 3.0 m/s. The negative slope represents leftward motion.",
            "misconception": "Students frequently plot the line starting at (0,0) or confuse a negative velocity with 'slowing down'."
        },
        {
            "title": "Problem 3: Multi-Stage Motion (Trot → Rest → Trot)",
            "calc": "• Stage 1: v₁ = (8.0 - 0) / 2.0 = +4.0 m/s\n• Stage 2: v₂ = (8.0 - 8.0) / 2.0 = 0.0 m/s (Rest)\n• Stage 3: v₃ = (14.0 - 8.0) / 2.0 = +3.0 m/s",
            "graph": "Piecewise horizontal steps: v = +4.0 m/s (0-2s) → v = 0.0 m/s on axis (2-4s) → v = +3.0 m/s (4-6s).",
            "map": "Evenly spaced pawprints (0 to 8m) → vertical cluster of pawprints at x = 8.0 m → slightly closer evenly spaced pawprints (8 to 14m).",
            "narrative": "Coco jogs fast right (+4 m/s), pauses to sniff for 2 seconds at x = 8m, then jogs right at a slightly slower pace (+3 m/s).",
            "misconception": "Students may think Stage 1 and Stage 3 have the same speed without calculating the slopes."
        },
        {
            "title": "Problem 4: Turnaround Motion (Fetch & Return)",
            "calc": "• Δx₁ = (+4.0)(2.0) = +8.0 m → x(2s) = 10.0 m\n• Δx₂ = 0.0 m → x(3s) = 10.0 m\n• Δx₃ = (-4.0)(2.0) = -8.0 m → x(5s) = 10.0 - 8.0 = 2.0 m\n• Total Distance = 8.0 + 0 + 8.0 = 16.0 m | Net Displacement = 2.0 - 2.0 = 0.0 m",
            "graph": "Triangular peak: Line from (0, 2m) to (2s, 10m) → flat line at 10m (2-3s) → line sloping down to (5s, 2m).",
            "map": "Rightward pawprints to 10m → stack at 10m → leftward pawprints back to 2m.",
            "narrative": "Coco runs right to fetch a ball, pauses at x = 10m, and returns to her starting position at x = 2m at the same speed.",
            "misconception": "Students often confuse net displacement (0 m) with total distance traveled (16 m)."
        }
    ]
    
    for sol in solutions:
        add_heading_2(doc, sol["title"])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run("• Mathematical Calculation: ").bold = True
        p.add_run(sol["calc"] + "\n")
        p.add_run("• Constructed Graph Solution: ").bold = True
        p.add_run(sol["graph"] + "\n")
        p.add_run("• Motion Map Feature: ").bold = True
        p.add_run(sol["map"] + "\n")
        p.add_run("• Written Interpretation: ").bold = True
        p.add_run(sol["narrative"] + "\n")
        p.add_run("• Key Student Misconception to Target: ").bold = True
        run_misc = p.add_run(sol["misconception"])
        run_misc.font.color.rgb = RGB_AMBER

    add_heading_1(doc, "🧠 Synthesis Questions Key")
    synth_ans = [
        "1. Horizontal Line Meaning: On a Position-Time graph, a horizontal line means position is not changing over time (Δx = 0, so v = 0, object is stopped). On a Velocity-Time graph, a horizontal line means velocity is not changing over time (velocity has a steady non-zero value, moving at constant speed).",
        "2. Distance vs Displacement: Total distance traveled is the actual path length covered (16.0 m). Net displacement is the overall straight-line change in position (x_final - x_initial = 2.0 m - 2.0 m = 0.0 m). Distance is a scalar, displacement is a vector.",
        "3. Simulation Alignment: Scenario CV-6 ('Fetch & Return') directly models Problem 4."
    ]
    for sa in synth_ans:
        p_sa = doc.add_paragraph()
        p_sa.paragraph_format.space_after = Pt(4)
        p_sa.add_run(sa)
        p_sa.runs[0].font.size = Pt(9)

    doc.save(filepath)
    print(f"Generated CV Teacher Guide: {filepath}")

if __name__ == "__main__":
    base_dir = "/Users/vladimir.lopez/Library/CloudStorage/GoogleDrive-vladimir.lopez@kinkaid.org/My Drive/AI/AI Workspace"
    student_out = os.path.join(base_dir, "motion_simulation", "Constant_Velocity_Motion_Handout.docx")
    teacher_out = os.path.join(base_dir, "motion_simulation", "Constant_Velocity_Motion_Teacher_Guide.docx")
    
    generate_cv_student_handout(student_out)
    generate_cv_teacher_guide(teacher_out)
