import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

HEX_TEAL = "0F7E9B"
HEX_AMBER = "D67B19"
HEX_LIGHT_BG = "E6F4F8"
HEX_LIGHT_AMBER = "FDF3E8"
HEX_DARK = "1E293B"

RGB_TEAL = RGBColor(15, 126, 155)
RGB_AMBER = RGBColor(214, 123, 25)
RGB_DARK = RGBColor(30, 41, 59)
RGB_MUTED = RGBColor(100, 116, 139)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13.5)
    run.font.bold = True
    run.font.color.rgb = RGB_TEAL
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGB_AMBER
    return p

def format_doc_header(doc, title, subtitle):
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("THE THINKING EXPERIMENT | KINEMATICS STUDIO")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(8.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGB_AMBER

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_TEAL

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(10)
    run_desc = p_desc.add_run(subtitle)
    run_desc.font.name = 'Arial'
    run_desc.font.size = Pt(9.5)
    run_desc.font.color.rgb = RGB_MUTED

def add_callout(doc, title, bullets, is_amber=False):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.8)
    cell = tbl.cell(0, 0)
    set_cell_background(cell, HEX_LIGHT_AMBER if is_amber else HEX_LIGHT_BG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run_t = p.add_run(title)
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10)
    run_t.font.bold = True
    run_t.font.color.rgb = RGB_AMBER if is_amber else RGB_TEAL
    
    for b in bullets:
        pb = cell.add_paragraph()
        pb.paragraph_format.space_after = Pt(2)
        run_b = pb.add_run(b)
        run_b.font.name = 'Arial'
        run_b.font.size = Pt(9)
        run_b.font.color.rgb = RGB_DARK
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)

def generate_student_handout(filepath):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
        
    format_doc_header(doc, "Kinematics Studio: Coco's Motion Lab", "Student Inquiry & Graphical Analysis Worksheet")
    
    add_callout(doc, "🎯 Learning Objectives — Constant Velocity (CVPM) & Acceleration (CAPM)", [
        "• Translate between physical dog motions, directional pawprint trails, Position-Time (x-t), and Velocity-Time (v-t) graphs.",
        "• Constant Velocity Rules: Slope of x-t = Velocity (v); Slope of v-t = 0 (horizontal line at v value).",
        "• Understand that flat horizontal lines on x-t mean stationary at rest (v = 0), while on v-t they mean constant speed.",
        "• Distinguish between linear piecewise segments (uniform motion) and parabolic curves (non-zero acceleration)."
    ])
    
    add_heading_1(doc, "🟢 Unit 1: Constant Velocity Particle Model (CVPM)")
    p_cv_intro = doc.add_paragraph("In the Constant Velocity tab, explore uniform motion where acceleration is zero (a = 0). Match Coco's motions to the correct graphs below:")
    p_cv_intro.paragraph_format.space_after = Pt(4)
    
    table_cv = doc.add_table(rows=9, cols=5)
    table_cv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cv.autofit = False
    
    headers = ["Card", "Type", "Motion Description (a = 0)", "Direction", "Speed Comparison"]
    col_widths = [Inches(0.8), Inches(1.1), Inches(2.7), Inches(0.9), Inches(1.3)]
    
    hdr_cells = table_cv.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], HEX_TEAL)
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=90, right=90)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    cv_data = [
        ("CV-1", "x-t", "Steady Rightward Trot: Starts at x=2m, trots right at +3.5 m/s.", "+ Right", "Uniform (+3.5 m/s)"),
        ("CV-2", "v-t", "Steady Leftward Trot: Starts at x=18m, trots left at -3.5 m/s.", "- Left", "Uniform (-3.5 m/s)"),
        ("CV-3", "x-t", "Patient Sniff: Sits at rest at x=12m sniffing a treat (v = 0).", "Rest (0)", "Zero Velocity"),
        ("CV-4", "v-t", "Fast Zoomie: Sprints right at high uniform velocity (+7.0 m/s).", "+ Right", "Fast Uniform"),
        ("CV-5", "x-t", "Trot, Sniff, Trot: Trots right (+4 m/s) → stops (v=0) → trots right (+4 m/s).", "+ Right", "Multi-Stage w/ Rest"),
        ("CV-6", "x-t", "Fetch & Return: Jogs right (+4.5 m/s) → grabs ball → runs back left (-4.5 m/s).", "+ then -", "Direction Reversal"),
        ("CV-7", "v-t", "Left Trot w/ Pause: Walks left (-3.5 m/s) → pauses (v=0) → walks left.", "- Left", "Multi-Stage (-v)"),
        ("CV-8", "x-t", "Speed Shift: Walks right at +2.5 m/s, then shifts to fast +6.0 m/s sprint.", "+ Right", "Shallow to Steep Slope")
    ]
    
    for row_idx, row_data in enumerate(cv_data, start=1):
        row_cells = table_cv.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = col_widths[col_idx]
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F8FAFC")
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGB_DARK
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGB_TEAL
                
    add_heading_1(doc, "🚀 Unit 2: Constant Acceleration Model (CAPM)")
    p_acc_intro = doc.add_paragraph("In the Full Kinematics tab, explore non-zero acceleration (a ≠ 0) involving curved Position graphs and sloping Velocity graphs:")
    p_acc_intro.paragraph_format.space_after = Pt(4)
    
    table_acc = doc.add_table(rows=11, cols=4)
    table_acc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_acc.autofit = False
    
    headers_acc = ["Mission", "Type", "Motion Description (a ≠ 0)", "Acceleration State"]
    col_widths_acc = [Inches(0.9), Inches(1.1), Inches(3.2), Inches(1.6)]
    
    hdr_cells_acc = table_acc.rows[0].cells
    for i, title in enumerate(headers_acc):
        hdr_cells_acc[i].width = col_widths_acc[i]
        set_cell_background(hdr_cells_acc[i], HEX_AMBER)
        set_cell_margins(hdr_cells_acc[i], top=80, bottom=80, left=90, right=90)
        p = hdr_cells_acc[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    acc_data = [
        ("#1", "x-t", "The Backyard Dash: Starts from rest at x=10m, speeds up right (+a).", "+a (Parabolic Curve)"),
        ("#2", "v-t", "The Sudden Treat Pause: Sprints left fast, smoothly decelerating to rest.", "+a Deceleration"),
        ("#3", "x-t", "The Boundary Patrol: Trots right steadily, then doubles speed right.", "Piecewise Constant"),
        ("#4", "v-t", "The Squirrel Hesitation: Slows to rest moving right, pauses, bolts left.", "Changing Sign (+ to -)"),
        ("#5", "x-t", "The Steady Sniff Tour: Trots left at const speed, stops to sniff, resumes left.", "Uniform w/ Rest Plateau"),
        ("#6", "x-t", "The Rebound Arc: Runs right slowing down, peaks, turns around left smoothly.", "Continuous Apex Reversal"),
        ("#7", "v-t", "Launch & Glide: Accelerates from rest right, then cruises at const speed.", "Acceleration to Cruise"),
        ("#8", "v-t", "Leftward Zoomies: Launches from rest at right, accelerating steadily left.", "-a (Negative Slope)"),
        ("#9", "x-t", "The Zig-Zag Search: Const speed right → turns const speed left → turns right.", "Piecewise 3-Stage"),
        ("#10", "v-t", "The Bell-Curve Trot: Accelerates right from rest → cruises → decelerates to rest.", "Trapezoidal Profile")
    ]
    
    for row_idx, row_data in enumerate(acc_data, start=1):
        row_cells = table_acc.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = col_widths_acc[col_idx]
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "FDF3E8")
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGB_DARK
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGB_AMBER

    doc.save(filepath)
    print(f"Generated Student Handout: {filepath}")

def generate_teacher_guide(filepath):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
        
    format_doc_header(doc, "Kinematics Studio: Coco's Motion Lab — Teacher Guide", "Complete Pedagogical Blueprint, Unit Taxonomy & Answer Keys")
    
    add_callout(doc, "📋 Unit Structuring for On-Level & Advanced Physics", [
        "• Unit 1 (CVPM / Constant Velocity): Focuses exclusively on uniform motion, initial position (y-intercept), direction of motion (sign of slope), and comparative speed (magnitude of slope). Zero acceleration throughout.",
        "• Unit 2 (CAPM / Acceleration): Introduces non-zero acceleration, curved position trajectories, linear velocity slopes, and turning points."
    ])
    
    add_heading_1(doc, "🔑 Complete Constant Velocity (CVPM) Answer Key")
    table_cv = doc.add_table(rows=9, cols=4)
    table_cv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cv.autofit = False
    
    headers_cv = ["Card", "Type", "Kinematic Analysis (a = 0)", "Targeted PER Misconception"]
    col_widths_cv = [Inches(0.8), Inches(0.9), Inches(2.8), Inches(2.3)]
    
    hdr_cells_cv = table_cv.rows[0].cells
    for i, title in enumerate(headers_cv):
        hdr_cells_cv[i].width = col_widths_cv[i]
        set_cell_background(hdr_cells_cv[i], HEX_TEAL)
        set_cell_margins(hdr_cells_cv[i], top=80, bottom=80, left=90, right=90)
        p = hdr_cells_cv[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    cv_answers = [
        ("CV-1", "x-t", "x₀=2m, straight line with positive slope v=+3.5 m/s.", "Starting at x=2m does not mean starting from rest."),
        ("CV-2", "v-t", "Horizontal line at v=-3.5 m/s below 0-axis for 4.0s.", "Confusing negative velocity with slowing down."),
        ("CV-3", "x-t", "Horizontal flat line at x=12m (slope = 0 m/s).", "Confusing flat position line with moving at constant speed."),
        ("CV-4", "v-t", "Horizontal line at v=+7.0 m/s above 0-axis for 2.5s.", "Relating vertical height on v-t to physical track position."),
        ("CV-5", "x-t", "Straight +slope → flat plateau (rest) → straight +slope.", "Confusing flat pause with turning backward."),
        ("CV-6", "x-t", "Straight +slope → flat pause → straight -slope back to x=2m.", "Thinking an inverted V-shape is walking up a physical hill."),
        ("CV-7", "v-t", "Line at -3.5 m/s → line on 0-axis → line at -3.5 m/s.", "Missing that 0 on v-t means stationary at rest."),
        ("CV-8", "x-t", "Shallow +slope (v=+2.5 m/s) connected to steep +slope (v=+6 m/s).", "Failing to connect line steepness to physical speed.")
    ]
    
    for row_idx, row_data in enumerate(cv_answers, start=1):
        row_cells = table_cv.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = col_widths_cv[col_idx]
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F8FAFC")
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGB_DARK
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGB_TEAL

    doc.save(filepath)
    print(f"Generated Teacher Guide: {filepath}")

if __name__ == "__main__":
    base_dir = "/Users/vladimir.lopez/Library/CloudStorage/GoogleDrive-vladimir.lopez@kinkaid.org/My Drive/AI/AI Workspace"
    student_out = os.path.join(base_dir, "motion_simulation", "Graph_That_Motion_Student_Handout.docx")
    teacher_out = os.path.join(base_dir, "motion_simulation", "Graph_That_Motion_Teacher_Guide.docx")
    
    generate_student_handout(student_out)
    generate_teacher_guide(teacher_out)
