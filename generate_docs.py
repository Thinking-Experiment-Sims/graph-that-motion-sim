import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Constants (PhysicsKit Palette) ---
HEX_TEAL = "0F7E9B"
HEX_AMBER = "D67B19"
HEX_DARK = "1E293B"
HEX_LIGHT_BG = "F8FAFC"
HEX_BORDER = "CBD5E1"
HEX_TEAL_LIGHT = "E6F4F8"

RGB_TEAL = RGBColor(15, 126, 155)
RGB_AMBER = RGBColor(214, 123, 25)
RGB_DARK = RGBColor(30, 41, 59)
RGB_MUTED = RGBColor(100, 116, 139)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name in ['top', 'left', 'bottom', 'right']:
        if b_name in kwargs:
            val, sz, col = kwargs[b_name]
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), val)
            node.set(qn('w:sz'), str(sz))
            node.set(qn('w:color'), col)
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_callout(doc, title_text, body_text_list, border_hex=HEX_TEAL, bg_hex=HEX_TEAL_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    set_cell_background(cell, bg_hex)
    set_cell_border(cell, left=('single', 24, border_hex))
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(title_text)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_TEAL if border_hex == HEX_TEAL else RGB_AMBER
    
    for item in body_text_list:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(item)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.color.rgb = RGB_DARK
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_doc_header(doc, title, subtitle):
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_TEAL
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGB_AMBER

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGB_TEAL
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGB_AMBER
    return p

# =========================================================================
# 1. BUILD STUDENT HANDOUT (.docx)
# =========================================================================
def generate_student_handout(output_path):
    doc = Document()
    
    # Page setup (0.75 in margins)
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        
    format_doc_header(doc, "Graph That Motion — Coco's Kinematics Adventure", "The Thinking Experiment | PhysicsKit Kinematics Module")
    
    # Callout Box: Objectives
    add_callout(doc, "🎯 Learning Objectives", [
        "• Translate between physical motion, ticker-tape motion maps (pawprint dots), Position-Time (x-t), and Velocity-Time (v-t) graphs.",
        "• Apply slope rules: Slope of x-t = Velocity (v); Slope of v-t = Acceleration (a).",
        "• Recognize that flat horizontal lines mean constant value (at rest on x-t; steady speed on v-t).",
        "• Relate curved parabolic shapes on x-t graphs to changing velocity (non-zero acceleration)."
    ])
    
    add_heading_1(doc, "🧭 Motion Map Interpretation Key")
    p_map = doc.add_paragraph()
    p_map.paragraph_format.space_after = Pt(8)
    p_map.add_run("Coco drops an amber pawprint dot at fixed time intervals (Δt = 0.30 s):\n").font.color.rgb = RGB_DARK
    p_map.add_run("• Constant Velocity: ").bold = True
    p_map.add_run("Dots are evenly spaced along the track.\n")
    p_map.add_run("• Speeding Up: ").bold = True
    p_map.add_run("Dots spread progressively further apart as Coco bolts forward.\n")
    p_map.add_run("• Slowing Down: ").bold = True
    p_map.add_run("Dots bunch progressively closer together as Coco brakes.\n")
    p_map.add_run("• Stationary (At Rest): ").bold = True
    p_map.add_run("Dots cluster and stack vertically while Coco sits or sniffs the ground.")
    
    add_heading_1(doc, "📝 Part 1: Graph Matching Reference Matrix")
    p_table_intro = doc.add_paragraph("Watch Coco's adventures in the simulation. Record the kinematics breakdown for each graph letter below:")
    p_table_intro.paragraph_format.space_after = Pt(6)
    
    # Table (12 rows, 5 cols)
    table = doc.add_table(rows=12, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Graph", "Type", "Coco's Motion Description", "Dir (+ / -)", "Speed State"]
    col_widths = [Inches(0.7), Inches(1.2), Inches(2.6), Inches(0.9), Inches(1.1)]
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], HEX_TEAL)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    matrix_data = [
        ("A", "Position-Time", "Coco starts from rest at left, accelerating right.", "+", "Accelerating"),
        ("B", "Velocity-Time", "Trots left at const speed → pauses to sniff → accelerates right.", "- then +", "Multi-Stage"),
        ("C", "Velocity-Time", "Trots steadily to the left at constant negative speed.", "-", "Constant"),
        ("D", "Position-Time", "Trots left slowing down → pauses to rest → continues trotting left.", "-", "Multi-Stage"),
        ("E", "Position-Time", "Trots smoothly right at constant speed from origin.", "+", "Constant"),
        ("F", "Velocity-Time", "Sprints left from rest → stops abruptly → walks left at const pace.", "-", "Multi-Stage"),
        ("G", "Position-Time", "Runs right for a ball → grabs it (rest) → trots back left.", "+ then -", "Turnaround"),
        ("H", "Velocity-Time", "Starts in a fast zoomie right, steadily braking to a stop.", "+", "Decelerating"),
        ("I", "Position-Time", "Spots squirrel left, accelerating from rest with increasing speed left.", "-", "Accelerating"),
        ("J", "Position-Time", "Walks right at steady pace → pauses → bolts forward in sprint.", "+", "Multi-Stage"),
        ("K", "Position-Time", "Dashes left fast, skidding with positive acceleration to a stop.", "-", "Decelerating")
    ]
    
    for row_idx, row_data in enumerate(matrix_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            top=('single', 4, HEX_BORDER), 
                            bottom=('single', 4, HEX_BORDER),
                            left=('single', 4, HEX_BORDER),
                            right=('single', 4, HEX_BORDER))
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(cell_value)
            r.font.name = 'Arial'
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGB_DARK
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGB_TEAL
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    add_heading_1(doc, "🔍 Part 2: Guided Conceptual Analysis")
    
    questions = [
        ("1. Slope vs. Position Comparison:", 
         "Compare Graph E and Graph A. Both start at position x = 0. How do their slopes differ, and what does that tell you about the car's speed over time?",
         ["Answer: ____________________________________________________________________________________",
          "____________________________________________________________________________________________"]),
        
        ("2. Visualizing Direction Changes:", 
         "How does a reversal of direction appear on a Position-Time graph (like Graph G) compared to a Velocity-Time graph (like Graph B)?",
         ["Position-Time (x-t): ________________________________________________________________________",
          "Velocity-Time (v-t): ________________________________________________________________________"]),
        
        ("3. Physical Meaning of Zero Velocity:", 
         "When an object is stationary at rest for 1.5 seconds, describe what you see across all 3 representations:",
         ["• Position-Time (x-t) shape: ________________________________________________________________",
          "• Velocity-Time (v-t) shape: ________________________________________________________________",
          "• Motion Map dots: ________________________________________________________________________"]),
        
        ("4. Acceleration Sign & Direction:", 
         "In Graph K, the car is moving in the negative direction (left), but the graph curves upward flattening out. Is the acceleration positive or negative? Explain using Δv.",
         ["Answer: ____________________________________________________________________________________",
          "____________________________________________________________________________________________"])
    ]
    
    for q_title, q_prompt, q_lines in questions:
        add_heading_2(doc, q_title)
        p_q = doc.add_paragraph(q_prompt)
        p_q.paragraph_format.space_after = Pt(4)
        for line in q_lines:
            p_l = doc.add_paragraph(line)
            p_l.paragraph_format.space_after = Pt(2)
            p_l.runs[0].font.color.rgb = RGB_MUTED
            p_l.runs[0].font.size = Pt(9.5)
            
    add_heading_1(doc, "💡 Part 3: Partner Challenge (Scenario Builder)")
    p_bld = doc.add_paragraph("Open the 🛠️ Scenario Builder tab in the simulation. Design a 2-stage motion, record your input parameters below, and challenge your desk partner to predict the resulting graphs!")
    p_bld.paragraph_format.space_after = Pt(6)
    
    add_callout(doc, "⚙️ Custom Challenge Design Parameters", [
        "• Stage 1:  Start Position x₀ = _______ m | Initial Velocity v₀ = _______ m/s | Accel a₁ = _______ m/s² | Time Δt₁ = _______ s",
        "• Stage 2:  Action (Continue / Rest / Reverse): ________________ | Accel a₂ = _______ m/s² | Time Δt₂ = _______ s",
        "• Partner Prediction Verification:  Did your partner correctly identify the graph shapes?  [  ] Yes   [  ] Needs Revision"
    ], border_hex=HEX_AMBER, bg_hex="FDF3E8")
    
    doc.save(output_path)
    print(f"Generated: {output_path}")

# =========================================================================
# 2. BUILD TEACHER GUIDE (.docx)
# =========================================================================
def generate_teacher_guide(output_path):
    doc = Document()
    
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        
    format_doc_header(doc, "Graph That Motion — Teacher Guide & Lesson Plan", "The Thinking Experiment | PhysicsKit Kinematics Module")
    
    add_callout(doc, "📋 Module Overview & Pedagogical Targets", [
        "This interactive module features Coco the Black Pug to target the 4 most persistent misconceptions in 1D kinematics:",
        "1. Graph-as-Picture Fallacy: Students mistaking graphical peaks/troughs for physical hills.",
        "2. Position vs. Velocity Confusion: Assuming high vertical position implies high speed.",
        "3. Slope vs. Height Confusion: Overlooking that steepness represents velocity (x-t) or acceleration (v-t).",
        "4. Sign of Acceleration Confusion: Believing negative acceleration always means 'slowing down'."
    ])
    
    add_heading_1(doc, "🔑 Complete Answer Key & Diagnostic Analysis")
    
    table = doc.add_table(rows=12, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Graph", "Type", "Coco's Kinematic Motion Breakdown", "Common Student Misconceptions"]
    col_widths = [Inches(0.7), Inches(1.1), Inches(2.6), Inches(2.1)]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], HEX_TEAL)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    ans_data = [
        ("A", "x-t", "Coco starts from rest (x₀=0, v₀=0, a=+2.7 m/s²). Parabolic curve steepening.", "Mistaking straight linear line for speeding up."),
        ("B", "v-t", "Coco trots left (v=-4 m/s) → sniffs at rest (v=0) → accelerates right.", "Confusing v = 0 axis line with 'at the origin'."),
        ("C", "v-t", "Coco trots left at constant steady velocity v = -4.2 m/s for 3.8s.", "Confusing negative velocity with deceleration."),
        ("D", "x-t", "Coco slows down moving left to stop → pauses → continues trotting left.", "Thinking curve flattening means moving right."),
        ("E", "x-t", "Coco trots right with steady constant velocity v = +4.2 m/s from origin.", "Baseline positive constant velocity reference."),
        ("F", "v-t", "Coco accelerates left (0 to -5.4 m/s) → pauses (v=0) → cruises left.", "Missing that negative slope on v-t is negative acceleration."),
        ("G", "x-t", "Coco runs right for ball → pauses to grab it → trots back left (peak).", "Thinking Coco ran up a triangular hill."),
        ("H", "v-t", "Coco starts in zoomie right (v=+8.5 m/s), braking linearly to stop.", "Confusing downward slope on v-t with moving backwards."),
        ("I", "x-t", "Coco spots squirrel, starts from rest at right, accelerating left.", "Confusing concave down curve with slowing down."),
        ("J", "x-t", "Coco walks right at const speed → pauses → bolts right in a sprint.", "Distinguishing linear segment from parabolic curve."),
        ("K", "x-t", "Coco dashes left fast (v₀=-8.5 m/s), skidding with +a to a stop.", "Recognizing that v < 0 and a > 0 produces deceleration.")
    ]
    
    for row_idx, row_data in enumerate(ans_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=90, right=90)
            set_cell_border(row_cells[col_idx], 
                            top=('single', 4, HEX_BORDER), 
                            bottom=('single', 4, HEX_BORDER),
                            left=('single', 4, HEX_BORDER),
                            right=('single', 4, HEX_BORDER))
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(cell_value)
            r.font.name = 'Arial'
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGB_DARK
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGB_TEAL
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    add_heading_1(doc, "🏫 4 Active Classroom Facilitation Strategies")
    
    strategies = [
        ("1. Four Corners Kinesthetic Match (15–20 min)", [
            "• Setup: Project the simulation at the front of the room. Label the 4 walls/corners with large letter signs (A, B, C, D).",
            "• Procedure: Play an animation (e.g. Motion #4). Give pairs 30 seconds to analyze the motion map dots, then call 'Commit!'. Students walk to their chosen graph corner.",
            "• Discourse: Call on one spokesperson from each occupied corner to defend their choice using physical evidence (e.g., 'The dots got closer together before stopping, which means slope flattens out...')."
        ]),
        ("2. Whiteboard Prediction Sketching (10–15 min)", [
            "• Setup: Distribute small whiteboards and markers to student pairs.",
            "• Procedure: Switch the simulation to the '✏️ Prediction Sketch' tab. Play the motion. Students must sketch their predicted x-t and v-t graphs before revealing any letter choices.",
            "• Pedagogical Benefit: Capitalizes on the 'Generation Effect'—forcing cognitive retrieval prior to multiple-choice recognition increases retention by over 30%."
        ]),
        ("3. Socratic Diagnostic Debriefs (Formative Interventions)", [
            "• Strategy: When a student makes an error, avoid giving the correct answer. Use the simulation's diagnostic prompts:",
            "  - 'Look at the vertical axis label. Does this graph tell us where the car is, or how fast it is moving?'",
            "  - 'Use the Time Scrubber to pause at t = 1.5s. Look at the amber velocity arrow. Which way is it pointing?'"
        ]),
        ("4. Scenario Builder Peer-Review Stations (15–20 min)", [
            "• Setup: Have students switch to the '🛠️ Scenario Builder' tab.",
            "• Task: Students program a custom 2-stage or 3-stage motion and challenge their partner to sketch the exact motion map and graph curves.",
            "• Extension: Ask advanced students to design a motion where velocity is negative but acceleration is positive, and explain the physical feeling of riding inside that car."
        ])
    ]
    
    for s_title, s_points in strategies:
        add_heading_2(doc, s_title)
        for pt in s_points:
            p = doc.add_paragraph(pt)
            p.paragraph_format.space_after = Pt(3)
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGB_DARK
            
    doc.save(output_path)
    print(f"Generated: {output_path}")

if __name__ == '__main__':
    base_dir = "/Users/vladimir.lopez/Library/CloudStorage/GoogleDrive-vladimir.lopez@kinkaid.org/My Drive/AI/AI Workspace/motion_simulation"
    generate_student_handout(os.path.join(base_dir, "Graph_That_Motion_Student_Handout.docx"))
    generate_teacher_guide(os.path.join(base_dir, "Graph_That_Motion_Teacher_Guide.docx"))
